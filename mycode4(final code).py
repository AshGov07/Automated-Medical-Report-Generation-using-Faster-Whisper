import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import threading
import os
import signal
import sys
import re
import time
import wave
import pyaudio
from datetime import datetime
from docx import Document
from RealtimeSTT import AudioToTextRecorder

class GynecologyReportUI:
    def __init__(self, root):
        self.root = root
        self.root.title("First Trimester Gynecology Report System")
        self.root.geometry("900x700")
        self.root.configure(bg="#f0f0f0")
        
        # Default save path - can be changed by user
        self.save_path = os.path.expanduser("~/Desktop/wav files")
        if not os.path.exists(self.save_path):
            os.makedirs(self.save_path)
        self.doc_path = os.path.join(self.save_path, "First_Trimester_Report.docx")
        
        # First Trimester Gynecology Report headings
        self.headings = [
            "Patient Information:",
            "LMP",
            "Gestational Age:",
            "Type of Scan:",
            "Uterine Position:",
            "Endometrial Thickness:",
            "Fetal Pole:",
            "Crown Rump Length",
            "Fetal Heart Rate:",
            "Amniotic Fluid:",
            "Placental Position:",
            "Adnexal Region:",
            "Cervical Length:",
            "Nuchal Translucency",
            "Additional Findings:",
            "Impression:",
            "Recommendations:"
        ]
        
        # Speech recognition variables
        self.current_heading = None
        self.command_buffer = []
        self.is_command_mode = False
        self.last_text_time = 0
        self.BUFFER_TIMEOUT = 3
        self.COMMAND_HISTORY_SIZE = 3
        self.is_recording = False
        self.recorder = None
        
        # Expanded command patterns to handle common misrecognitions
        go_to_patterns = [
            r'^go to (.+)$',
            r'^goto (.+)$',
            r'^go do (.+)$',
            r'^go do it (.+)$',
            r'^go 2 (.+)$',
            r'^go too (.+)$',
            r'^go toward (.+)$',
            r'^go through (.+)$',
            r'^go (.+)$',
            r'^to (.+)$'
        ]
        self.command_patterns = [re.compile(pattern) for pattern in go_to_patterns]
        
        # Audio recording variables
        self.audio = None
        self.stream = None
        self.frames = []
        self.is_audio_recording = False
        
        # Create UI elements
        self.create_menu()
        self.create_ui()
        
        # Initialize Word document if it doesn't exist
        self.init_document()
        
        # Initialize buffer checker thread
        self.buffer_thread = threading.Thread(target=self.check_buffer, daemon=True)
        self.buffer_thread.start()

    def create_menu(self):
        menubar = tk.Menu(self.root)
        
        # File menu
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="New Report", command=self.new_report)
        file_menu.add_command(label="Open Report", command=self.open_report)
        file_menu.add_command(label="Change Save Location", command=self.change_save_location)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.exit_app)
        menubar.add_cascade(label="File", menu=file_menu)
        
        # Help menu
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="Instructions", command=self.show_instructions)
        help_menu.add_command(label="About", command=self.show_about)
        menubar.add_cascade(label="Help", menu=help_menu)
        
        self.root.config(menu=menubar)

    def create_ui(self):
        # Main frame
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Top section - Physician and controls
        top_frame = ttk.Frame(main_frame)
        top_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(top_frame, text="Physician Name:").pack(side=tk.LEFT, padx=5)
        self.physician_var = tk.StringVar()
        physician_entry = ttk.Entry(top_frame, textvariable=self.physician_var, width=30)
        physician_entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(top_frame, text="Update", command=self.update_physician).pack(side=tk.LEFT, padx=5)
        
        # Recording control buttons
        self.record_button = ttk.Button(top_frame, text="Start Recording", command=self.toggle_recording)
        self.record_button.pack(side=tk.RIGHT, padx=5)
        
        # Middle section - Split view
        middle_frame = ttk.Frame(main_frame)
        middle_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # Left side - Headings list
        left_frame = ttk.LabelFrame(middle_frame, text="Report Sections")
        left_frame.pack(side=tk.LEFT, fill=tk.Y, padx=5, expand=False)
        
        self.heading_listbox = tk.Listbox(left_frame, width=30, height=20)
        self.heading_listbox.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        for heading in self.headings:
            self.heading_listbox.insert(tk.END, heading)
        self.heading_listbox.bind("<<ListboxSelect>>", self.on_heading_select)
        
        # Right side - Transcription and content
        right_frame = ttk.Frame(middle_frame)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)
        
        # Currently selected section
        section_frame = ttk.Frame(right_frame)
        section_frame.pack(fill=tk.X, pady=5)
        ttk.Label(section_frame, text="Current Section:").pack(side=tk.LEFT, padx=5)
        self.current_section_var = tk.StringVar(value="None selected")
        ttk.Label(section_frame, textvariable=self.current_section_var, font=("", 10, "bold")).pack(side=tk.LEFT, padx=5)
        
        # Content box for current section
        content_frame = ttk.LabelFrame(right_frame, text="Section Content")
        content_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.content_text = scrolledtext.ScrolledText(content_frame, wrap=tk.WORD, width=50, height=10)
        self.content_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Transcription log
        log_frame = ttk.LabelFrame(right_frame, text="Transcription Log")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.log_text = scrolledtext.ScrolledText(log_frame, wrap=tk.WORD, width=50, height=10)
        self.log_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.log_text.config(state=tk.DISABLED)
        
        # Bottom section - Status bar
        status_frame = ttk.Frame(main_frame)
        status_frame.pack(fill=tk.X, pady=5)
        
        self.status_var = tk.StringVar(value="Ready. Select a section and start recording.")
        status_bar = ttk.Label(status_frame, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(fill=tk.X)

    def init_document(self):
        # Initialize Word document if it doesn't exist
        if not os.path.exists(self.doc_path):
            self.log("Creating new report document...")
            doc = Document()
            doc.add_heading('First Trimester Ultrasound Report', level=1)
            
            # Add date and time
            current_datetime = datetime.now().strftime("%B %d, %Y at %I:%M %p")
            doc.add_paragraph(f"Report Date: {current_datetime}")
            
            # Add doctor information placeholder
            doc.add_paragraph("Physician: Dr. _________________")
            doc.add_paragraph("--------------------------------------------------")
            
            for heading in self.headings:
                doc.add_paragraph(heading, style='Heading 2')
                doc.add_paragraph("")  # Empty paragraph for text input
            
            # Add signature section
            doc.add_paragraph("--------------------------------------------------")
            doc.add_paragraph("Signature: _________________")
            doc.add_paragraph("Date: _________________")
            
            doc.save(self.doc_path)
            self.log("New report document created successfully.")
        else:
            self.log("Existing report document found.")

    def normalize_text(self, text):
        # Normalize text (remove punctuation and make lowercase)
        return re.sub(r'[^\w\s]', '', text).strip().lower()

    def update_word_document(self, heading, text):
        try:
            doc = Document(self.doc_path)
            found = False
            
            # First check for actual heading paragraphs (style='Heading 2')
            for i in range(len(doc.paragraphs) - 1):
                if doc.paragraphs[i].style.name == 'Heading 2' and self.normalize_text(doc.paragraphs[i].text) == self.normalize_text(heading):
                    doc.paragraphs[i + 1].text = text  # Update the paragraph after the heading
                    found = True
                    break
            
            # If not found by style, try with text content (backwards compatibility)
            if not found:
                for i in range(len(doc.paragraphs) - 1):
                    if self.normalize_text(doc.paragraphs[i].text) == self.normalize_text(heading):
                        doc.paragraphs[i + 1].text = text  # Update the paragraph after the heading
                        found = True
                        break
                        
            if found:
                doc.save(self.doc_path)
                self.log(f"Updated: {heading} -> {text}")
                
                # Update the content text widget to show the current value
                if heading == self.current_heading:
                    self.content_text.delete(1.0, tk.END)
                    self.content_text.insert(tk.END, text)
            else:
                self.log(f"Error: Heading '{heading}' not found!")
        except Exception as e:
            self.log(f"Error updating document: {e}")

    def is_potential_command_start(self, text):
        normalized = self.normalize_text(text)
        
        # Common misrecognitions of "go to"
        command_starters = ['go to', 'goto', 'go do', 'go do it', 'go 2', 'go too', 'go toward', 'go through', 'go', 'to']
        
        for starter in command_starters:
            if normalized.startswith(starter):
                return True
        
        return False

    def check_for_command_in_history(self, text_history):
        # Join the history fragments in various combinations to catch commands split across fragments
        possible_texts = []
        
        # Add the individual fragments
        possible_texts.extend(text_history)
        
        # Try joining adjacent fragments
        for i in range(len(text_history)-1):
            possible_texts.append(f"{text_history[i]} {text_history[i+1]}")
        
        # Try joining all fragments
        if len(text_history) > 1:
            possible_texts.append(" ".join(text_history))
        
        # Check each possibility against our command patterns
        for text in possible_texts:
            normalized = self.normalize_text(text)
            for pattern in self.command_patterns:
                match = pattern.match(normalized)
                if match:
                    return match, pattern
        
        return None, None

    def process_text(self, text):
        if not text.strip():
            return  # Skip empty text
        
        current_time = time.time()
        
        # Add text to log
        self.log(f"Received text: '{text}' | Command mode: {self.is_command_mode}")
        
        # Add to command buffer for history-based detection
        self.command_buffer.append(text)
        if len(self.command_buffer) > self.COMMAND_HISTORY_SIZE:
            self.command_buffer.pop(0)
        
        # First, check if the current text is a direct command
        normalized_text = self.normalize_text(text)
        command_match = None
        matched_pattern = None
        
        for pattern in self.command_patterns:
            match = pattern.match(normalized_text)
            if match:
                command_match = match
                matched_pattern = pattern
                break
        
        # If direct command found, process it
        if command_match:
            target_heading = command_match.group(1).strip()
            self.log(f"Command detected: {matched_pattern.pattern} -> '{target_heading}'")
            
            # Find matching heading
            for heading in self.headings:
                if self.normalize_text(heading).startswith(target_heading):
                    self.current_heading = heading
                    # Update UI to show current heading
                    self.current_section_var.set(heading)
                    # Select in listbox
                    for i, h in enumerate(self.headings):
                        if h == heading:
                            self.heading_listbox.selection_clear(0, tk.END)
                            self.heading_listbox.selection_set(i)
                            self.heading_listbox.see(i)
                            # Load content for this heading
                            self.load_current_heading_content()
                            break
                    
                    self.log(f"Switched to: {self.current_heading}")
                    self.is_command_mode = False
                    self.command_buffer = []
                    return
            
            self.log(f"No matching heading found for: {target_heading}")
            self.is_command_mode = False
            return
        
        # Second, check the command buffer for fragmented commands
        command_match, matched_pattern = self.check_for_command_in_history(self.command_buffer)
        if command_match:
            target_heading = command_match.group(1).strip()
            self.log(f"Command detected in history: {matched_pattern.pattern} -> '{target_heading}'")
            
            # Find matching heading
            for heading in self.headings:
                if self.normalize_text(heading).startswith(target_heading):
                    self.current_heading = heading
                    # Update UI to show current heading
                    self.current_section_var.set(heading)
                    # Select in listbox
                    for i, h in enumerate(self.headings):
                        if h == heading:
                            self.heading_listbox.selection_clear(0, tk.END)
                            self.heading_listbox.selection_set(i)
                            self.heading_listbox.see(i)
                            # Load content for this heading
                            self.load_current_heading_content()
                            break
                    
                    self.log(f"Switched to: {self.current_heading}")
                    self.is_command_mode = False
                    self.command_buffer = []
                    return
            
            self.log(f"No matching heading found for: {target_heading}")
            self.is_command_mode = False
            self.command_buffer = []
            return
        
        # Third, check if text might be the start of a command
        if self.is_potential_command_start(text):
            self.is_command_mode = True
            self.last_text_time = current_time
            self.log(f"Potential command detected: {text} (buffering...)")
            return
        
        # If we're in command mode, don't add text to document until we verify it's not a command
        if self.is_command_mode:
            self.last_text_time = current_time
            self.log(f"Still in command mode, buffering...")
            return
        
        # IMPORTANT: Additional safety check - if text contains common command triggers, don't add to document
        command_triggers = ['go to', 'goto', 'go do', 'go']
        contains_trigger = False
        for trigger in command_triggers:
            if trigger in normalized_text:
                contains_trigger = True
                self.log(f"Text contains command trigger '{trigger}', not adding to document")
                # Enter command mode to process potential command
                self.is_command_mode = True
                self.last_text_time = current_time
                return
        
        # Regular text and we have a current heading - update document
        if self.current_heading and not contains_trigger:
            self.log(f"Adding text to {self.current_heading}: {text}")
            self.update_word_document(self.current_heading, text)
        elif not self.current_heading:
            self.log(f"No heading selected. Say 'go to [heading]' or select one from the list.")
            self.status_var.set("Please select a section first!")

    def check_buffer(self):
        while True:
            time.sleep(0.5)
            current_time = time.time()
            
            # If we're in command mode and buffer timeout has elapsed
            if self.is_command_mode and (current_time - self.last_text_time) > self.BUFFER_TIMEOUT:
                self.log(f"Command buffer timeout - exiting command mode")
                self.is_command_mode = False
                # Don't clear command_buffer - keep history for future commands

    def update_physician(self):
        name = self.physician_var.get().strip()
        if not name:
            messagebox.showwarning("Warning", "Please enter a physician name.")
            return
            
        try:
            doc = Document(self.doc_path)
            # Look for the physician line (typically the third line)
            for i, para in enumerate(doc.paragraphs):
                if para.text.startswith("Physician:"):
                    doc.paragraphs[i].text = f"Physician: Dr. {name}"
                    doc.save(self.doc_path)
                    self.log(f"Updated physician name to Dr. {name}")
                    messagebox.showinfo("Success", f"Physician name updated to Dr. {name}")
                    return
            
            self.log(f"Physician line not found in document.")
            messagebox.showwarning("Warning", "Physician line not found in document.")
        except Exception as e:
            self.log(f"Error updating physician name: {e}")
            messagebox.showerror("Error", f"Error updating physician name: {e}")

    def toggle_recording(self):
        if not self.is_recording:
            self.start_recording()
        else:
            self.stop_recording()

    def start_recording(self):
        # Start audio recording in a separate thread
        self.is_recording = True
        self.record_button.config(text="Stop Recording")
        self.status_var.set("Recording... Speak clearly")
        
        # Configure STT recorder
        recorder_config = {
            'spinner': False,
            'model': 'large-v2',
            'language': 'en',
            'silero_sensitivity': 0.4,
            'webrtc_sensitivity': 2,
            'post_speech_silence_duration': 0.6,
            'min_length_of_recording': 0,
            'min_gap_between_recordings': 0,
            'enable_realtime_transcription': True,
            'realtime_processing_pause': 0.3,
            'realtime_model_type': 'tiny.en',
            'on_realtime_transcription_update': self.process_text,
        }
        
        self.recorder = AudioToTextRecorder(**recorder_config)
        
        # Start audio recording for wav file
        self.audio_thread = threading.Thread(target=self.record_audio)
        self.audio_thread.start()
        
        # Start transcription
        self.transcription_thread = threading.Thread(target=self.run_transcription)
        self.transcription_thread.start()
        
        self.log("Recording and transcription started...")

    def stop_recording(self):
        self.is_recording = False
        self.record_button.config(text="Start Recording")
        self.status_var.set("Recording stopped")
        
        # Stop STT recorder
        if self.recorder:
            self.recorder.stop()
        
        # Stop audio recording
        self.is_audio_recording = False
        
        self.log("Recording and transcription stopped")

    def record_audio(self):
        # Function to record audio to a wav file
        file_name = os.path.join(self.save_path, "recorded_audio.wav")
        self.audio = pyaudio.PyAudio()
        self.stream = self.audio.open(format=pyaudio.paInt16, channels=1, rate=44100, input=True, frames_per_buffer=1024)
        self.frames = []
        self.is_audio_recording = True
        
        try:
            while self.is_audio_recording:
                data = self.stream.read(1024)
                self.frames.append(data)
        except Exception as e:
            self.log(f"Error recording audio: {e}")
        
        # Save recorded audio
        self.stream.stop_stream()
        self.stream.close()
        self.audio.terminate()
        
        wf = wave.open(file_name, 'wb')
        wf.setnchannels(1)
        wf.setsampwidth(self.audio.get_sample_size(pyaudio.paInt16))
        wf.setframerate(44100)
        wf.writeframes(b''.join(self.frames))
        wf.close()
        
        self.log(f"Audio saved to {file_name}")

    def run_transcription(self):
        # Run the transcription in a separate thread
        try:
            while self.is_recording:
                self.recorder.text(self.process_text)
                time.sleep(0.1)
        except Exception as e:
            self.log(f"Error in transcription: {e}")

    def on_heading_select(self, event):
        selection = self.heading_listbox.curselection()
        if selection:
            index = selection[0]
            self.current_heading = self.headings[index]
            self.current_section_var.set(self.current_heading)
            self.log(f"Selected heading: {self.current_heading}")
            self.load_current_heading_content()

    def load_current_heading_content(self):
        # Load content for current heading from document
        try:
            doc = Document(self.doc_path)
            found = False
            
            # First check for headings by style
            for i in range(len(doc.paragraphs) - 1):
                if doc.paragraphs[i].style.name == 'Heading 2' and self.normalize_text(doc.paragraphs[i].text) == self.normalize_text(self.current_heading):
                    content = doc.paragraphs[i + 1].text
                    self.content_text.delete(1.0, tk.END)
                    self.content_text.insert(tk.END, content)
                    found = True
                    break
            
            # If not found by style, try with text content
            if not found:
                for i in range(len(doc.paragraphs) - 1):
                    if self.normalize_text(doc.paragraphs[i].text) == self.normalize_text(self.current_heading):
                        content = doc.paragraphs[i + 1].text
                        self.content_text.delete(1.0, tk.END)
                        self.content_text.insert(tk.END, content)
                        found = True
                        break
                        
            if not found:
                self.log(f"Content for heading '{self.current_heading}' not found!")
                self.content_text.delete(1.0, tk.END)
        except Exception as e:
            self.log(f"Error loading content: {e}")

    def log(self, message):
        # Add message to log with timestamp
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_message = f"[{timestamp}] {message}\n"
        
        # Update log text widget
        self.log_text.config(state=tk.NORMAL)
        self.log_text.insert(tk.END, log_message)
        self.log_text.see(tk.END)  # Scroll to bottom
        self.log_text.config(state=tk.DISABLED)
        
        # For debugging
        print(log_message)

    # Menu functions
    def new_report(self):
        if os.path.exists(self.doc_path):
            if messagebox.askyesno("Confirm", "This will create a new report document. Are you sure?"):
                # Backup existing file
                backup_path = f"{self.doc_path}.bak"
                os.rename(self.doc_path, backup_path)
                self.log(f"Existing report backed up to {backup_path}")
                
                # Create new document
                self.init_document()
                self.content_text.delete(1.0, tk.END)
                messagebox.showinfo("Success", "New report created")
        else:
            self.init_document()
            messagebox.showinfo("Success", "New report created")

    def open_report(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Word Documents", "*.docx")],
            initialdir=self.save_path
        )
        if file_path:
            self.doc_path = file_path
            self.log(f"Opened report: {file_path}")
            self.status_var.set(f"Working with: {os.path.basename(file_path)}")
            
            # Reset current heading
            self.current_heading = None
            self.current_section_var.set("None selected")
            self.content_text.delete(1.0, tk.END)

    def change_save_location(self):
        directory = filedialog.askdirectory(initialdir=self.save_path)
        if directory:
            self.save_path = directory
            self.doc_path = os.path.join(self.save_path, "First_Trimester_Report.docx")
            self.log(f"Save location changed to: {directory}")
            self.status_var.set(f"Save location: {directory}")

    def show_instructions(self):
        instructions = """
Voice Commands:
- Say "go to [section name]" to switch to a different section
- Recognized variations: "go to", "goto", "go do", etc.
- Example: "go to patient information"

Tips:
- Select a section before starting recording
- Update physician name before starting
- Speak clearly and pause between phrases
- Text will be automatically added to the current section
"""
        messagebox.showinfo("Instructions", instructions)

    def show_about(self):
        about_text = """
First Trimester Gynecology Report System
Version 1.0

This application allows voice-controlled creation of first
trimester gynecology reports with automatic transcription.

© 2025 All rights reserved
"""
        messagebox.showinfo("About", about_text)

    def exit_app(self):
        if messagebox.askyesno("Exit", "Are you sure you want to exit?"):
            if self.is_recording:
                self.stop_recording()
            self.root.destroy()
            sys.exit(0)

def main():
    root = tk.Tk()
    app = GynecologyReportUI(root)
    root.protocol("WM_DELETE_WINDOW", app.exit_app)
    root.mainloop()

if __name__ == "__main__":
    main()